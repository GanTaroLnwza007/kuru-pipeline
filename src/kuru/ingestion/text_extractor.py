"""Text extractor — PyMuPDF for born-digital PDFs, per-page Typhoon for image pages.

Bulk OCR (for fully scanned PDFs) lives in ocr_extractor.py and is NOT called here.
To re-enable scanned OCR:
    from kuru.ingestion.ocr_extractor import extract_with_ocr
"""

from __future__ import annotations

import base64
import os
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF

from kuru.ingestion.utils import png_dimensions, safe_print

# If PyMuPDF extracts fewer than this many chars total, the PDF is treated as scanned.
SCANNED_CHAR_THRESHOLD = 500

# Pages with fewer than this many chars get routed to per-page Typhoon OCR.
_PAGE_LOW_YIELD_CHARS = 50

_PAGE_OCR_DPI = 96


@dataclass
class PageText:
    page_num: int
    text: str
    extraction_method: str  # 'pymupdf' | 'typhoon_page' | 'scanned' | 'python-docx' | 'failed'


# ─────────────────────────────────────────
# Internal helpers (exported for testing)
# ─────────────────────────────────────────

def _should_ocr_page(text: str) -> bool:
    """True when a page's extracted text is too short to be useful."""
    return len(text.strip()) < _PAGE_LOW_YIELD_CHARS


def _extract_page_typhoon(page_b64: str) -> str:
    """Send a single low-yield page to Typhoon OCR. Returns '' if unavailable."""
    if not os.environ.get("TYPHOON_API_KEY"):
        return ""
    try:
        from kuru.llm import get_ocr_client  # noqa: PLC0415

        w, h = png_dimensions(page_b64)
        prompt = (
            f"Below is an image of a document page. The image dimensions are {w}x{h} pixels. "
            "Extract all text from this page. Preserve Thai text exactly as written. "
            "Output only the extracted text with no commentary."
        )
        response = get_ocr_client().chat.completions.create(
            model="typhoon-ocr",
            messages=[{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{page_b64}"}},
            ]}],
            temperature=0.1,
            max_tokens=4096,
            top_p=0.6,
            extra_body={"repetition_penalty": 1.2},
        )
        if not response.choices:
            return ""
        return response.choices[0].message.content or ""
    except Exception:
        return ""


def _build_page_texts(pdf_path: Path, verbose: bool = False) -> list[PageText]:
    """Extract text page-by-page. Low-yield pages are sent to per-page Typhoon."""
    doc = fitz.open(str(pdf_path))
    pages: list[PageText] = []

    for i, page in enumerate(doc):
        text = page.get_text("text")
        method = "pymupdf"

        if _should_ocr_page(text):
            b64 = base64.b64encode(
                page.get_pixmap(dpi=_PAGE_OCR_DPI).tobytes("png")
            ).decode()
            ocr_text = _extract_page_typhoon(b64)
            if ocr_text.strip():
                text = ocr_text
                method = "typhoon_page"
                if verbose:
                    safe_print(f"  [typhoon] ✓ page {i + 1} ({len(text)} chars)")

        pages.append(PageText(page_num=i, text=text, extraction_method=method))

    doc.close()
    return pages


# ─────────────────────────────────────────
# Public API
# ─────────────────────────────────────────

def extract_text(
    pdf_path: str | Path,
    use_vision_fallback: bool = True,  # kept for call-site compatibility, ignored
    verbose: bool = False,
) -> list[PageText]:
    """Extract text from a PDF using PyMuPDF + optional per-page Typhoon for image pages.

    If total extracted chars < SCANNED_CHAR_THRESHOLD, all pages are marked 'scanned'
    and no OCR is attempted. Re-enable full OCR by importing extract_with_ocr from
    kuru.ingestion.ocr_extractor.
    """
    pdf_path = Path(pdf_path)
    pages = _build_page_texts(pdf_path, verbose=verbose)
    total_chars = sum(len(p.text.strip()) for p in pages)

    if total_chars < SCANNED_CHAR_THRESHOLD:
        if verbose:
            safe_print(
                f"  Low text yield ({total_chars} chars) — marking as scanned. "
                "OCR disabled; import ocr_extractor.extract_with_ocr to re-enable."
            )
        for p in pages:
            p.extraction_method = "scanned"

    return pages


def render_page_b64(pdf_path: Path, page_num: int, dpi: int = 150) -> str:
    """Render a single PDF page to a base64 PNG string."""
    doc = fitz.open(str(pdf_path))
    pix = doc[page_num].get_pixmap(dpi=dpi)
    doc.close()
    return base64.b64encode(pix.tobytes("png")).decode()


def extract_text_from_docx(docx_path: Path) -> list[PageText]:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx  # python-docx  # noqa: PLC0415

        doc = docx.Document(str(docx_path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append("\t".join(row_cells))
        return [PageText(page_num=0, text="\n".join(parts), extraction_method="python-docx")]
    except Exception as exc:
        safe_print(f"  DOCX extraction failed ({type(exc).__name__}): {exc}")
        return [PageText(page_num=0, text="", extraction_method="failed")]


def extract_text_auto(
    path: str | Path,
    use_vision_fallback: bool = True,
    verbose: bool = False,
) -> list[PageText]:
    """Dispatch to the right extractor based on file extension (.pdf or .docx)."""
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return extract_text_from_docx(path)
    return extract_text(path, use_vision_fallback=use_vision_fallback, verbose=verbose)


def full_text(pages: list[PageText]) -> str:
    return "\n\n".join(p.text for p in pages if p.text.strip())
